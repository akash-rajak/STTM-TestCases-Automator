
'''

##'2suvzridx73mx7vpr5jpvciy76hgxmfe2ffbxay2mtf35tw6654a'
##'https://dev.azure.com/akashr0568/'
## 'STTM_Automation'

'''

## imported necessary library
import pandas as pd
import openai
import xlwt
from xlwt import Workbook
import openpyxl
import os
import pathlib
from pathlib import Path
from tkinter import filedialog, Tk
from tkinter.filedialog import askdirectory
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.styles.alignment import Alignment
import docx
import json

from azure.devops.connection import Connection
from msrest.authentication import BasicAuthentication
from azure.devops.v6_0.work_item_tracking.models import JsonPatchOperation, WorkItem

## api key defined
openai.api_key = "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"

ado_info = {
    'ado_flag': 0,
    'pat': "",
    'url': '',
    'project': '',
    'parent_id': -1

}
## connection defined
ado_flag = 0
personal_access_token = ''
organization_url = ''


# credentials = BasicAuthentication('', personal_access_token)
# connection = Connection(base_url=organization_url, creds=credentials)


## function defined to read excel file
def readExcel(file):
    df = pd.read_excel(file, sheet_name="Requirement")

    x = -1
    for i in range(len(df[df.columns[0]])):
        if "High Level Requirements" in str(df[df.columns[0]][i]):
            x = i

    sheet_names = []
    for i in df.columns[1:]:
        sheet_names.append(i)
    if (x != -1):
        s = df.iloc[x]
        s = s[1:]
        for i in s:
            sheet_names.append(i)

    start = -1

    wb = openpyxl.Workbook()
    t_c = []
    for sheet in range(len(sheet_names)):
        data = pd.read_excel(file, sheet_name=sheet_names[sheet])
        s = -1
        t = -1
        # print(data.columns)
        for i in range(len(data.columns)):
            temp = data.columns[i]
            temp = str(temp)
            temp = temp.lower()
            if "source" in temp and (s == -1):
                s = i
            elif "target" in temp and (t == -1):
                t = i
        # print(s)
        # print(t)

        start = -1
        for i in range(len(data)):
            if (start == -1):
                for j in range(len(data.iloc[i])):
                    temp = data.iloc[i, j]
                    temp = str(temp)
                    temp = temp.lower()
                    # print(temp)
                    if "source" in temp or "target" in temp or "transformation" in temp:
                        start = i
                        break
            else:
                break

        data_store = {
            "Source System": -1,
            "Source Server": -1,
            "Source DB": -1,
            "Source Table": -1,
            "Source Column": -1,
            "Source Data Type": -1,
            "Source Primary Key": -1,
            "Source Allowed Nulls": -1,
            "Target System": -1,
            "Target Server": -1,
            "Target DB": -1,
            "Target Table": -1,
            "Target Column": -1,
            "Target Data Type": -1,
            "Target Primary Key": -1,
            "Target Allowed Nulls": -1,
            "Data Transformation Rules": -1,

        }

        for i in range(len(data.iloc[start])):
            temp = data.iloc[start, i]
            temp = str(temp)
            if "Transformation" in str(temp):
                data_store['Data Transformation Rules'] = i
            elif (i >= t) or "Target" in temp:
                if "Target System" in str(temp):
                    data_store['Target System'] = i
                elif "Target Server" in str(temp):
                    data_store['Target Server'] = i
                elif "Target DB" in str(temp) or "Target Database" in str(temp):
                    data_store['Target DB'] = i
                elif "Target File" in str(temp) or "Target Table" in str(temp):
                    data_store["Target Table"] = i
                elif "Target Column" in str(temp):
                    data_store['Target Column'] = i
                elif "data" in str(temp).lower() or "target data type" in str(temp).lower() or "data type" in str(
                        temp).lower():
                    data_store['Target Data Type'] = i
                elif "primary key" in str(temp).lower() or "key" in str(temp).lower() or "target primarykey" in str(
                        temp).lower():
                    data_store['Target Primary Key'] = i
                elif "nulls" in str(temp).lower():
                    data_store['Target Allowed Nulls'] = i
            else:
                if "Source System" in str(temp):
                    data_store['Source System'] = i
                elif "Source Server" in str(temp):
                    data_store['Source Server'] = i
                elif "Source DB" in str(temp) or "Source Database" in str(temp):
                    data_store['Source DB'] = i
                elif "Source File" in str(temp) or "Source  File" in str(temp) or "Source Table" in str(temp):
                    data_store["Source Table"] = i
                elif "Source Column" in str(temp):
                    data_store['Source Column'] = i
                elif "data" in str(temp).lower() or "source data type" in str(temp).lower() or "data type" in str(
                        temp).lower():
                    data_store['Source Data Type'] = i
                elif "primary key" in str(temp).lower() or "key" in str(temp).lower() or "source primarykey" in str(
                        temp).lower():
                    data_store['Source Primary Key'] = i
                elif "nulls" in str(temp).lower():
                    data_store['Source Allowed Nulls'] = i
        tc = []
        tct = 0
        tt = 0
        start += 1
        for i in range(start, len(data)):
            # print(i)
            source_table = data.iloc[i, data_store["Source Table"]]
            source_column = data.iloc[i, data_store["Source Column"]]
            source_datatype = data.iloc[i, data_store["Source Data Type"]]

            rule = data.iloc[i, data_store['Data Transformation Rules']]

            target_table = data.iloc[i, data_store["Target Table"]]
            target_column = data.iloc[i, data_store["Target Column"]]
            target_datatype = data.iloc[i, data_store["Target Data Type"]]
            prompt = f"""
            Generate test cases queries both for source and target and validation for source to target mapping with data integration, 
            source to target mapping, data validation, data count, datatype checking and casting if source and target data types are different
            based on the rule {rule}:
            Source Table = {source_table}
            Source Column = {source_column}
            Source Datatype = {source_datatype}
            Target Table = {target_table}
            Target Column = {target_column}
            Target Datatype = {target_datatype}
            Generate output in tabular format Test Case No.|Test Case Type|Source Query|Target Query|Validation Query|
            """
            # print(prompt)
            response = openai.Completion.create(
                engine="text-davinci-003",
                prompt=prompt,
                max_tokens=512,
                temperature=0.2,
                seed=29
            )
            source_Information = {
                "Source System": data.iloc[i, data_store["Source System"]],
                "Source Server": data.iloc[i, data_store["Source Server"]],
                "Source Database": data.iloc[i, data_store["Source DB"]],
                "Source Table": data.iloc[i, data_store["Source Table"]],
                "Source Column": data.iloc[i, data_store["Source Column"]],
                "Source DataType": data.iloc[i, data_store["Source Data Type"]],
                "Primary Key": data.iloc[i, data_store["Source Primary Key"]],
                "Allowed Nulls": data.iloc[i, data_store["Source Allowed Nulls"]],
            }
            target_Information = {
                "Target System": data.iloc[i, data_store["Target System"]],
                "Target Server": data.iloc[i, data_store["Target Server"]],
                "Target Database": data.iloc[i, data_store["Target DB"]],
                "Target Table": data.iloc[i, data_store["Target Table"]],
                "Target Column": data.iloc[i, data_store["Target Column"]],
                "Target DataType": data.iloc[i, data_store["Target Data Type"]],
                "Primary Key": data.iloc[i, data_store["Target Primary Key"]],
                "Allowed Nulls": data.iloc[i, data_store["Target Allowed Nulls"]],
            }
            # response
            a = response.choices[0].text.strip()
            # print(a)
            prompt = f"""
            Generate for each query generated Test Case Type, Test Case Number, Test Case Description, Test Case Steps, for the given queries generated {a}
            output should be in the format:
            source information: {source_Information}
            target information: {target_Information}
            Test Case Number|Test Case Type|Test Case Name|Test Case Description|Test Case Steps|Source Query|Target Query|Validation Query|Expected Output|
            And the first column 'Test Case Number' should contain only digits without including leading zeroes and Use '\n' only if you need to switch to a new test case, otherwise u can use a 'tab' if needed.
            For Example: First row of the output looks like this, 1|Validate Data|Validate Data type of CountryCode for the Source and Target|This test case validates the data type of the CountryCode for source and target| "Step 1: Validate the data type of CountryCode for source and target Step 2: Check the syntax of the query"|SELECT CountryCode FROM fdn_sales_public|SELECT CountryCode from fdn_ibp.MacroEconomic|SELECT CountryCode from fdn_ibp.MacroEconomic|The source data type should be same as the target one.
            """
            responses = openai.Completion.create(
                engine="text-davinci-003",
                prompt=prompt,
                max_tokens=768,
                seed=10206,
                temperature=0.5
            )
            test_cases = responses.choices[0].text.strip()
            # print(test_cases)
            tc.append(test_cases)
        t_c.append(tc)
        dir_name = os.path.dirname(file)
        new_dir = pathlib.Path(dir_name, "Txt Output")
        new_dir.mkdir(parents=True, exist_ok=True)
        file_name = "TestCases_" + sheet_names[sheet] + ".txt"
        save1 = str(new_dir) + "\\" + file_name
        print("Created " + file_name)
        with open(save1, 'w', encoding='utf-8') as f:
            f.write("\n----------------------------------------------------------------------------\n")
            # f.write("\n\n\n\n\n")
            # f.write("Test Case No. |\t| Test case Type |\t| Test Case Description |\t| Source Database |\t| Source Table |\t| Source Column |\t| Target Database |\t| Target Table |\t| Target Column|\t| Source Query |\t| Target Query |\t| Expected Output \n")
            for i in range(len(tc)):
                q = tc[i].split("\n")
                for k in q:
                    if (
                            "Test Case Number | Test Case Type " in k
                            or "Test Case Number  |Test Case Type " in k
                            or "Test Case Number|Test Case Type" in k
                            or "____________________" in k
                            or "-----------------" in k
                            or "Test Case No. | Test Case Type" in k
                            or "Test Case Number ||| Test Case Type " in k
                            or "Test Case No. |Test Case Type" in k
                            or "Test Case Number |Test Case Type" in k
                            or "Test Case Number|Test Case Type" in k
                            or ":-----" in k
                            or "| Test Case Number | Test Case Type " in k
                            or "Test Case Number  |  Test Case Type " in k
                    ):
                        # print(k)
                        pass
                    else:
                        k = k.replace("<br>", "\n")
                        k = k.replace("\t", "\n")
                        si = f"""
                        Source System: {data.iloc[i + start, data_store['Source System']] if (data_store['Source System'] != -1) else "NA"}\nSource Server: {data.iloc[i + start, data_store['Source Server']] if (data_store['Source Server'] != -1) else "NA"}\nSource Database: {data.iloc[i + start, data_store['Source DB']] if (data_store['Source DB'] != -1) else "NA"}\nSource Table: {data.iloc[i + start, data_store['Source Table']] if (data_store['Source Table'] != -1) else "NA"}\nSource Column: {data.iloc[i + start, data_store['Source Column']] if (data_store['Source Column'] != -1) else "NA"}\nSource DataType: {data.iloc[i + start, data_store['Source Data Type']] if (data_store['Source Data Type'] != -1) else "NA"}\nPrimary Key: {data.iloc[i + start, data_store['Source Primary Key']] if (data_store['Source Primary Key'] != -1) else "NA"}\nAllowed Nulls: {data.iloc[i + start, data_store['Source Allowed Nulls']] if (data_store['Source Allowed Nulls'] != -1) else "NA"}
                        """
                        ti = f"""
                        Target System: {data.iloc[i + start, data_store['Target System']] if (data_store['Target System'] != -1) else "NA"}\nTarget Server: {data.iloc[i + start, data_store['Target Server']] if (data_store['Target Server'] != -1) else "NA"}\nTarget Database: {data.iloc[i + start, data_store['Target DB']] if (data_store['Target DB'] != -1) else "NA"}\nTarget Table: {data.iloc[i + start, data_store['Target Table']] if (data_store['Target Table'] != -1) else "NA"}\nTarget Column: {data.iloc[i + start, data_store['Target Column']] if (data_store['Target Column'] != -1) else "NA"}\nTarget DataType: {data.iloc[i + start, data_store['Target Data Type']] if (data_store['Target Data Type'] != -1) else "NA"}\nPrimary Key: {data.iloc[i + start, data_store['Target Primary Key']] if (data_store['Target Primary Key'] != -1) else "NA"}\nAllowed Nulls: {data.iloc[i + start, data_store['Target Allowed Nulls']] if (data_store['Target Allowed Nulls'] != -1) else "NA"}
                        """
                        f.write(si)
                        f.write('\n')
                        f.write(ti)
                        f.write('\n')
                        k.replace('|', '\n')
                        f.write(k)
                        f.write("\n")
        dir_name = os.path.dirname(file)
        new_dir = pathlib.Path(dir_name, "Word Output")
        new_dir.mkdir(parents=True, exist_ok=True)
        file_name = "TestCases_" + sheet_names[sheet] + ".docx"
        print("Created " + file_name)
        save1 = str(new_dir) + "\\" + file_name
        doc = docx.Document()
        doc.add_paragraph("\n----------------------------------------------------------------------------\n")
        para = doc.paragraphs[0]
        for i in range(len(tc)):
            q = tc[i].split("\n")
            # print(q)
            for k in q:
                # print(k)
                if (
                        "Test Case Number | Test Case Type " in k
                        or "Test Case Number  |Test Case Type " in k
                        or "Test Case Number|Test Case Type" in k
                        or "____________________" in k
                        or "-----------------" in k
                        or "Test Case No. | Test Case Type" in k
                        or "Test Case Number ||| Test Case Type " in k
                        or "Test Case No. |Test Case Type" in k
                        or "Test Case Number |Test Case Type" in k
                        or "Test Case Number|Test Case Type" in k
                        or ":-----" in k
                        or "| Test Case Number | Test Case Type " in k
                        or "Test Case Number  |  Test Case Type " in k
                ):
                    # print(k)
                    pass
                else:
                    si = f"""
                    Source System: {data.iloc[i + start, data_store['Source System']] if (data_store['Source System'] != -1) else "NA"}\nSource Server: {data.iloc[i + start, data_store['Source Server']] if (data_store['Source Server'] != -1) else "NA"}\nSource Database: {data.iloc[i + start, data_store['Source DB']] if (data_store['Source DB'] != -1) else "NA"}\nSource Table: {data.iloc[i + start, data_store['Source Table']] if (data_store['Source Table'] != -1) else "NA"}\nSource Column: {data.iloc[i + start, data_store['Source Column']] if (data_store['Source Column'] != -1) else "NA"}\nSource DataType: {data.iloc[i + start, data_store['Source Data Type']] if (data_store['Source Data Type'] != -1) else "NA"}\nPrimary Key: {data.iloc[i + start, data_store['Source Primary Key']] if (data_store['Source Primary Key'] != -1) else "NA"}\nAllowed Nulls: {data.iloc[i + start, data_store['Source Allowed Nulls']] if (data_store['Source Allowed Nulls'] != -1) else "NA"}
                    """
                    ti = f"""
                    Target System: {data.iloc[i + start, data_store['Target System']] if (data_store['Target System'] != -1) else "NA"}\nTarget Server: {data.iloc[i + start, data_store['Target Server']] if (data_store['Target Server'] != -1) else "NA"}\nTarget Database: {data.iloc[i + start, data_store['Target DB']] if (data_store['Target DB'] != -1) else "NA"}\nTarget Table: {data.iloc[i + start, data_store['Target Table']] if (data_store['Target Table'] != -1) else "NA"}\nTarget Column: {data.iloc[i + start, data_store['Target Column']] if (data_store['Target Column'] != -1) else "NA"}\nTarget DataType: {data.iloc[i + start, data_store['Target Data Type']] if (data_store['Target Data Type'] != -1) else "NA"}\nPrimary Key: {data.iloc[i + start, data_store['Target Primary Key']] if (data_store['Target Primary Key'] != -1) else "NA"}\nAllowed Nulls: {data.iloc[i + start, data_store['Target Allowed Nulls']] if (data_store['Target Allowed Nulls'] != -1) else "NA"}
                    """
                    # x=k.split('|')
                    k = k.replace('|', '\n')
                    # print(k)
                    para.add_run(si.strip())
                    para.add_run('\n')
                    para.add_run(ti.strip())
                    para.add_run('\n')
                    # for x in k:
                    # if(len(x.strip())>=1):
                    para.add_run(k.strip())
                    para.add_run('\n')

                para.add_run("\n\n\n")
        doc.save(save1)
        if (ado_info['ado_flag'] == 1):
            credentials = BasicAuthentication('', ado_info['pat'])
            connection = Connection(base_url=ado_info['url'], creds=credentials)
        wb = openpyxl.Workbook()
        heading = [
            'Test Case Number', 'Test Case Type', 'Test Case Name',
            'Test Case Description', 'Test Case Steps', 'Source Information',
            'Target Information', 'Source Query', 'Target Query', 'Validation Query', 'Expected Output'
        ]
        ch = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K']

        ws = wb.create_sheet("Test Cases Sheet " + str(sheet + 1))
        for i in range(len(heading)):
            cell = ws.cell(1, i + 1)
            cell.value = heading[i]
            ws.column_dimensions[ch[i]].width = 30

        ws.column_dimensions['D'].width = 50
        ws.column_dimensions['E'].width = 50
        ws.column_dimensions['H'].width = 50
        ws.column_dimensions['I'].width = 50
        ws.column_dimensions['J'].width = 50
        cnt = 2
        # flag=0
        for i in range(len(tc)):
            # =i.replace('|\n','|\n|')
            q = tc[i].split('\n')
            cell = ws.cell(cnt, 1)
            cell.value = "New Column" + str(i + 1)
            cnt += 1
            for k in q:
                if (
                        "Test Case Number | Test Case Type " in k
                        or "Test Case Number  |Test Case Type " in k
                        or "Test Case Number|Test Case Type" in k
                        or "____________________" in k
                        or "-----------------" in k
                        or "Test Case No. | Test Case Type" in k
                        or "Test Case Number ||| Test Case Type " in k
                        or "Test Case No. |Test Case Type" in k
                        or "Test Case Number |Test Case Type" in k
                        or "Test Case Number|Test Case Type" in k
                        or ":-----" in k
                        or "| Test Case Number | Test Case Type " in k
                        or "Test Case Number  |  Test Case Type " in k
                ):
                    # print(k)
                    pass
                else:
                    cnt1 = 1
                    k = k.replace("\t", "\n")
                    k = k.replace("<br>", "\n")
                    k = k.replace("<\br>", "\n")
                    if '|' not in k:
                        x = k.split("\n")
                    else:
                        x = k.split('|')
                    # print("x :", x)
                    si = f"""
                    Source System: {data.iloc[i + start, data_store['Source System']] if (data_store['Source System'] != -1) else "NA"}\nSource Server: {data.iloc[i + start, data_store['Source Server']] if (data_store['Source Server'] != -1) else "NA"}\nSource Database: {data.iloc[i + start, data_store['Source DB']] if (data_store['Source DB'] != -1) else "NA"}\nSource Table: {data.iloc[i + start, data_store['Source Table']] if (data_store['Source Table'] != -1) else "NA"}\nSource Column: {data.iloc[i + start, data_store['Source Column']] if (data_store['Source Column'] != -1) else "NA"}\nSource DataType: {data.iloc[i + start, data_store['Source Data Type']] if (data_store['Source Data Type'] != -1) else "NA"}\nPrimary Key: {data.iloc[i + start, data_store['Source Primary Key']] if (data_store['Source Primary Key'] != -1) else "NA"}\nAllowed Nulls: {data.iloc[i + start, data_store['Source Allowed Nulls']] if (data_store['Source Allowed Nulls'] != -1) else "NA"}
                    """
                    ti = f"""
                    Target System: {data.iloc[i + start, data_store['Target System']] if (data_store['Target System'] != -1) else "NA"}\nTarget Server: {data.iloc[i + start, data_store['Target Server']] if (data_store['Target Server'] != -1) else "NA"}\nTarget Database: {data.iloc[i + start, data_store['Target DB']] if (data_store['Target DB'] != -1) else "NA"}\nTarget Table: {data.iloc[i + start, data_store['Target Table']] if (data_store['Target Table'] != -1) else "NA"}\nTarget Column: {data.iloc[i + start, data_store['Target Column']] if (data_store['Target Column'] != -1) else "NA"}\nTarget DataType: {data.iloc[i + start, data_store['Target Data Type']] if (data_store['Target Data Type'] != -1) else "NA"}\nPrimary Key: {data.iloc[i + start, data_store['Target Primary Key']] if (data_store['Target Primary Key'] != -1) else "NA"}\nAllowed Nulls: {data.iloc[i + start, data_store['Target Allowed Nulls']] if (data_store['Target Allowed Nulls'] != -1) else "NA"}
                    """
                    source_db = data.iloc[i + start, data_store['Source DB']] if (
                                data_store['Source DB'] != -1) else "NA"
                    target_db = data.iloc[i + start, data_store['Target DB']] if (
                                data_store['Target DB'] != -1) else "NA"
                    source_query = ""
                    target_query = ""
                    expected_output = ""
                    validation_query = ""
                    tc_name = "TestCase"
                    si.replace("\t", "")
                    ti.replace("\t", "")
                    cell = ws.cell(cnt, 6)
                    cell.value = si.strip()
                    cell = ws.cell(cnt, 7)
                    cell.value = ti.strip()
                    for j in x:
                        if (len(j.strip()) >= 1):
                            if (cnt1 == 6):
                                cnt1 += 2

                            cell = ws.cell(cnt, cnt1)
                            cell.value = j.strip()
                            if (cnt1 == 3):
                                tc_name = j.strip()
                            if (cnt1 == 8):
                                source_query = j.strip()
                            if (cnt1 == 9):
                                target_query = j.strip()
                            if (cnt1 == 10):
                                validation_query = j.strip()
                            if (cnt1 == 11):
                                expected_output = j.strip()
                            cnt1 += 1
                    # print(ado_flag)
                    if (ado_info['ado_flag'] == 1):

                        # print(connection.get_client)
                        testcase_steps = f"""<steps id="0" last="6"><step id="2" type="ValidateStep"><parameterizedString isformatted="true">&lt;DIV&gt;&lt;P&gt;Connect To Source Database: {source_db}&lt;/P&gt;&lt;/DIV&gt;</parameterizedString><parameterizedString isformatted="true">&lt;DIV&gt;&lt;P&gt;Source Database should be connected&lt;/P&gt;&lt;/DIV&gt;</parameterizedString><description/></step><step id="3" type="ValidateStep"><parameterizedString isformatted="true">&lt;DIV&gt;&lt;P&gt;Connect To Target Database: {target_db}&lt;/P&gt;&lt;/DIV&gt;</parameterizedString><parameterizedString isformatted="true">&lt;DIV&gt;&lt;P&gt;Target Database should be connected&lt;/P&gt;&lt;/DIV&gt;</parameterizedString><description/></step><step id="4" type="ValidateStep"><parameterizedString isformatted="true">&lt;DIV&gt;&lt;P&gt;Run this query for source: {source_query}&lt;/P&gt;&lt;/DIV&gt;</parameterizedString><parameterizedString isformatted="true">&lt;DIV&gt;&lt;P&gt;source information&lt;/P&gt;&lt;/DIV&gt;</parameterizedString><description/></step><step id="5" type="ValidateStep"><parameterizedString isformatted="true">&lt;DIV&gt;&lt;P&gt;Run this query for target: {target_query}&lt;/P&gt;&lt;/DIV&gt;</parameterizedString><parameterizedString isformatted="true">&lt;DIV&gt;&lt;P&gt;target information&lt;/P&gt;&lt;/DIV&gt;</parameterizedString><description/></step><step id="6" type="ValidateStep"><parameterizedString isformatted="true">&lt;DIV&gt;&lt;P&gt;Run this query: {validation_query}&lt;/P&gt;&lt;/DIV&gt;</parameterizedString><parameterizedString isformatted="true">&lt;DIV&gt;&lt;P&gt;{expected_output}&lt;/P&gt;&lt;/DIV&gt;</parameterizedString><description/></step></steps>"""
                        # print(testcase_steps)
                        wit_client = connection.clients.get_work_item_tracking_client()
                        parent_work_item_id = ado_info['parent_id']
                        project = ado_info['project']
                        task = 'Test Case'
                        new_work_item_fields = {
                            'System.Title': tc_name,
                            'System.Steps': testcase_steps,
                            'System.State': 'Design'
                        }
                        json_patch_operations = [
                            JsonPatchOperation(
                                op='add',
                                path='/fields/System.Title',
                                value=new_work_item_fields['System.Title']
                            ),

                            JsonPatchOperation(
                                op='add',
                                path='/fields/Microsoft.VSTS.TCM.Steps',
                                value=new_work_item_fields['System.Steps']
                            ),
                            JsonPatchOperation(
                                op='add',
                                path='/fields/System.State',
                                value=new_work_item_fields['System.State']
                            )

                        ]
                        organization_url = ado_info['url']
                        w_flag = 0
                        try:
                            wit_client = connection.clients.get_work_item_tracking_client()
                            work_item = wit_client.get_work_item(parent_work_item_id)
                            w_flag = 1
                        except:
                            w_flag = 0
                        if (parent_work_item_id != -1 and w_flag == 1):
                            json_patch_operations.append(
                                JsonPatchOperation(
                                    op='add',
                                    path='/relations/-',
                                    value={
                                        'rel': 'System.LinkTypes.Hierarchy-Reverse',
                                        'url': f'{organization_url}/{project}/_apis/wit/workItems/{parent_work_item_id}',
                                        'attributes': {
                                            'comment': 'Making a new link for the dependency',
                                            'name': 'Parent'
                                        }
                                    }
                                )
                            )
                        else:
                            print("Parent id not entered or does not exists")
                        response = wit_client.create_work_item(
                            project=project,
                            type=task,
                            document=json_patch_operations
                        )
                        # print(response.id)
                    cnt += 1

        table = Table(displayName="Testcases_Table_" + str(sheet + 1), ref="A1:K" + str(cnt - 1))
        ws.add_table(table)
        # Apply some style to the table
        style = TableStyleInfo(name="TableStyleMedium9", showFirstColumn=False, showLastColumn=False,
                               showRowStripes=True,
                               showColumnStripes=False)
        table.tableStyleInfo = style

        for col in ws.columns:
            for cell in col:
                cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

    dir_name = os.path.dirname(file)
    new_dir = pathlib.Path(dir_name, "EXCEL Output")
    new_dir.mkdir(parents=True, exist_ok=True)
    file_name = "TestCases_" + os.path.basename(file)
    print("Created " + file_name)
    save1 = str(new_dir) + "\\" + file_name
    wb.save(save1)
    wb = openpyxl.load_workbook(save1)
    sheet_to_remove = wb['Sheet']
    wb.remove(sheet_to_remove)
    wb.save(save1)


## function defined to get api from user
def api():
    key = input("Enter your OpenAI API Secret key : ")
    if (len(key) == len(openai.api_key)):
        openai.api_key = key


## created ADO function
def ADO_testcase():
    ado = input("Do you want to create test cases on ADO Dashboard(Y/N) : ")
    if ado == 'Y':
        organization_url = input("Enter Organization URL of Azure Devops : ")
        personal_access_token = input("Enter Personal Access Token : ")
        project_name = input("Enter project where you want to create test cases : ")
        p_id = int(input("Enter parent id : "))
        ado_flag = 1
        ado_info['ado_flag'] = ado_flag
        ado_info['pat'] = personal_access_token
        ado_info['url'] = organization_url
        ado_info['project'] = project_name
        ado_info['parent_id'] = p_id

    else:
        ado_flag = 0


## definition of main function
def main():
    print("Select STTM Excel File : ", end='')
    root = Tk()
    root.attributes("-topmost", True)
    root.withdraw()
    file = filedialog.askopenfilename(title="Select file")
    print(file)
    # print(os.path.basename(file))
    if (file != ""):
        print("\nCurrently Processing {" + str(Path(file).stem) + ".xlsx}...")
        readExcel(file)
    else:
        print("No File Selected")


api()
ADO_testcase()
main()